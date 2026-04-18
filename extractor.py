import av
import cv2
import numpy as np
import pytesseract

import json
import logging
import subprocess
from pathlib import Path

import pandas as pd
from PIL import Image, ImageEnhance, ImageFilter
from pypdf import PdfReader
import pdfplumber
from docx import Document
from bs4 import BeautifulSoup
from striprtf.striprtf import rtf_to_text

logger = logging.getLogger(__name__)

CSV_CHUNK_SIZE = 50_000
CSV_MAX_CHARS = 5_000_000
JSON_MAX_ITEMS = 10_000
PARQUET_MAX_ROWS = 100_000
PDF_MAX_PAGES = 50
EXCEL_MAX_ROWS = 50_000
IMAGE_MAX_SIDE = 2000


class TextExtractor:
    def extract(self, file_path: str) -> str:
        raise NotImplementedError


class CSVExtractor(TextExtractor):
    def extract(self, file_path: str) -> str:
        try:
            chunks = []
            total = 0
            for chunk in pd.read_csv(
                file_path,
                chunksize=CSV_CHUNK_SIZE,
                on_bad_lines='skip',
                low_memory=False,
                dtype=str
            ):
                s = chunk.to_string(index=False, header=False)
                chunks.append(s)
                total += len(s)
                if total > CSV_MAX_CHARS:
                    break
            return "\n".join(chunks)
        except Exception as e:
            logger.error(f"CSV read error {file_path}: {e}")
            return ""


class JSONExtractor(TextExtractor):
    def extract(self, file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                data = json.load(f)
            if isinstance(data, list):
                data = data[:JSON_MAX_ITEMS]
            return json.dumps(data, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.error(f"JSON read error {file_path}: {e}")
            return ""


class ParquetExtractor(TextExtractor):
    def extract(self, file_path: str) -> str:
        try:
            import pyarrow.parquet as pq
            pf = pq.ParquetFile(file_path)
            rows_left = PARQUET_MAX_ROWS
            parts = []
            for batch in pf.iter_batches(batch_size=10_000):
                if rows_left <= 0:
                    break
                df = batch.to_pandas()
                if len(df) > rows_left:
                    df = df.head(rows_left)
                parts.append(df.to_string(index=False))
                rows_left -= len(df)
            return "\n".join(parts)
        except Exception as e:
            logger.error(f"Parquet read error {file_path}: {e}")
            return ""


class PDFExtractor(TextExtractor):
    def extract(self, file_path: str) -> str:
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages[:PDF_MAX_PAGES]:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            logger.debug(f"pdfplumber failed {file_path}: {e}")

        if not text.strip():
            try:
                reader = PdfReader(file_path)
                for page in reader.pages[:PDF_MAX_PAGES]:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            except Exception as e:
                logger.error(f"pypdf failed {file_path}: {e}")
        return text


class DOCXExtractor(TextExtractor):
    def extract(self, file_path: str) -> str:
        try:
            doc = Document(file_path)
            parts = [p.text for p in doc.paragraphs if p.text]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            parts.append(cell.text)
            return "\n".join(parts)
        except Exception as e:
            logger.error(f"DOCX read error {file_path}: {e}")
            return ""


class DOCExtractor(TextExtractor):
    def extract(self, file_path: str) -> str:
        for tool in ('antiword', 'catdoc'):
            try:
                result = subprocess.run(
                    [tool, file_path],
                    capture_output=True,
                    timeout=60,
                    check=False,
                )
                if result.returncode == 0 and result.stdout:
                    return result.stdout.decode('utf-8', errors='ignore')
            except (FileNotFoundError, subprocess.TimeoutExpired) as e:
                logger.debug(f"{tool} not available or timed out: {e}")
                continue
        logger.warning(f"DOC read failed {file_path}: install antiword or catdoc")
        return ""


class RTFExtractor(TextExtractor):
    def extract(self, file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return rtf_to_text(f.read())
        except Exception as e:
            logger.error(f"RTF read error {file_path}: {e}")
            return ""


class XLSXExtractor(TextExtractor):
    def extract(self, file_path: str) -> str:
        try:
            sheets = pd.read_excel(
                file_path,
                engine='openpyxl',
                sheet_name=None,
                nrows=EXCEL_MAX_ROWS,
                dtype=str,
            )
            return "\n".join(
                df.to_string(index=False) for df in sheets.values()
            )
        except Exception as e:
            logger.error(f"XLSX read error {file_path}: {e}")
            return ""


class XLSExtractor(TextExtractor):
    def extract(self, file_path: str) -> str:
        try:
            sheets = pd.read_excel(
                file_path,
                engine='xlrd',
                sheet_name=None,
                nrows=EXCEL_MAX_ROWS,
                dtype=str,
            )
            return "\n".join(
                df.to_string(index=False) for df in sheets.values()
            )
        except Exception as e:
            logger.error(f"XLS read error {file_path}: {e}")
            return ""


class HTMLExtractor(TextExtractor):
    def extract(self, file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                soup = BeautifulSoup(f, 'html.parser')
            for tag in soup(['script', 'style']):
                tag.decompose()
            return soup.get_text(separator=' ')
        except Exception as e:
            logger.error(f"HTML read error {file_path}: {e}")
            return ""


def preprocess_image(img: Image.Image) -> Image.Image:
    img = ImageEnhance.Contrast(img).enhance(2.0)
    img = img.convert('L')
    img = img.filter(ImageFilter.MedianFilter(size=3))
    img = img.point(lambda x: 0 if x < 128 else 255, '1')
    return img


class ImageExtractor(TextExtractor):
    def __init__(self, use_ocr: bool = False):
        self.use_ocr = use_ocr

    def extract(self, file_path: str) -> str:
        if not self.use_ocr:
            return ""
        try:
            img = Image.open(file_path)
            img = preprocess_image(img)
            img.thumbnail((IMAGE_MAX_SIDE, IMAGE_MAX_SIDE), Image.Resampling.LANCZOS)
            return pytesseract.image_to_string(img, lang='rus+eng')
        except Exception as e:
            logger.error(f"OCR error {file_path}: {e}")
            return ""


class TXTExtractor(TextExtractor):
    MAX_SIZE = 5_000_000

    def extract(self, file_path: str) -> str:
        try:
            size = Path(file_path).stat().st_size
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read(min(size, self.MAX_SIZE))
        except Exception as e:
            logger.error(f"TXT read error {file_path}: {e}")
            return ""


class VideoExtractor(TextExtractor):
    MAX_FRAMES = 300

    def __init__(self, use_ocr=True, frame_interval_sec=1.0):
        self.use_ocr = use_ocr
        self.frame_interval_sec = frame_interval_sec

    def _detect_document_contour(self, image_np):
        try:
            gray = cv2.cvtColor(image_np, cv2.COLOR_BGR2GRAY)
            blur = cv2.GaussianBlur(gray, (5, 5), 0)
            edged = cv2.Canny(blur, 75, 200)
            contours, _ = cv2.findContours(edged, cv2.RETR_LIST, cv2.CHAIN_APPROX_SIMPLE)
            contours = sorted(contours, key=cv2.contourArea, reverse=True)[:5]
            for c in contours:
                peri = cv2.arcLength(c, True)
                approx = cv2.approxPolyDP(c, 0.02 * peri, True)
                if len(approx) == 4:
                    warped = self._four_point_transform(image_np, approx.reshape(4, 2))
                    return warped
            return None
        except Exception as e:
            logger.debug(f"Document detection failed: {e}")
            return None

    def _four_point_transform(self, image, pts):
        rect = self._order_points(pts)
        (tl, tr, br, bl) = rect
        widthA = np.linalg.norm(br - bl)
        widthB = np.linalg.norm(tr - tl)
        maxWidth = max(int(widthA), int(widthB))
        heightA = np.linalg.norm(tr - br)
        heightB = np.linalg.norm(tl - bl)
        maxHeight = max(int(heightA), int(heightB))
        dst = np.array([
            [0, 0],
            [maxWidth - 1, 0],
            [maxWidth - 1, maxHeight - 1],
            [0, maxHeight - 1]], dtype="float32")
        M = cv2.getPerspectiveTransform(rect, dst)
        warped = cv2.warpPerspective(image, M, (maxWidth, maxHeight))
        return warped

    def _order_points(self, pts):
        rect = np.zeros((4, 2), dtype="float32")
        s = pts.sum(axis=1)
        rect[0] = pts[np.argmin(s)]
        rect[2] = pts[np.argmax(s)]
        diff = np.diff(pts, axis=1)
        rect[1] = pts[np.argmin(diff)]
        rect[3] = pts[np.argmax(diff)]
        return rect

    def extract(self, file_path: str) -> str:
        if not self.use_ocr:
            return ""
        all_text = []
        try:
            container = av.open(file_path)
            video_stream = container.streams.video[0]
            fps = float(video_stream.average_rate)
            frame_interval = max(1, int(fps * self.frame_interval_sec))
            frame_count = 0
            processed = 0
            for frame in container.decode(video_stream):
                if processed >= self.MAX_FRAMES:
                    break
                if frame_count % frame_interval == 0:
                    img = frame.to_ndarray(format='bgr24')
                    doc_img = self._detect_document_contour(img)
                    if doc_img is None:
                        doc_img = img
                    pil_img = Image.fromarray(cv2.cvtColor(doc_img, cv2.COLOR_BGR2RGB))
                    text = pytesseract.image_to_string(pil_img, lang='rus+eng')
                    if text.strip():
                        all_text.append(text.strip())
                    processed += 1
                frame_count += 1
        except Exception as e:
            logger.error(f"Video processing error {file_path}: {e}")
        return "\n".join(all_text)


_EXTRACTORS = {
    '.csv': CSVExtractor,
    '.json': JSONExtractor,
    '.parquet': ParquetExtractor,
    '.pdf': PDFExtractor,
    '.docx': DOCXExtractor,
    '.doc': DOCExtractor,
    '.rtf': RTFExtractor,
    '.xlsx': XLSXExtractor,
    '.xls': XLSExtractor,
    '.htm': HTMLExtractor,
    '.html': HTMLExtractor,
    '.mp4': VideoExtractor,
    '.txt': TXTExtractor,
    '.md': TXTExtractor,
}

_IMAGE_EXTS = {'.tif', '.tiff', '.jpeg', '.jpg', '.png', '.gif'}


def get_extractor(file_path: str, use_ocr: bool = False):
    ext = Path(file_path).suffix.lower()
    if ext in _IMAGE_EXTS:
        return ImageExtractor(use_ocr)
    if ext == '.mp4':
        return VideoExtractor(use_ocr=use_ocr)
    cls = _EXTRACTORS.get(ext)
    return cls() if cls else None
